"""
DOCX generator for resume and cover letter — matches user's formatting style.

Generates .docx files with proper formatting:
  - Cover letter: Arial, 9.5pt body, date header, company address, 4-paragraph structure
  - Resume: Arial, 13pt name, 11pt section headings with bottom borders,
    9.5pt body, bold job titles with right-aligned dates, bullet points

Optional PDF conversion via LibreOffice when available.
"""

import logging
import re
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, Any

from docx import Document
from docx.shared import Pt, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

logger = logging.getLogger(__name__)

# ── Style Constants (from user's reference document) ─────────────────────

FONT_NAME = "Arial"
NAME_SIZE = Pt(13)       # 13pt for name heading
SECTION_SIZE = Pt(11)    # 11pt for section headings (PROFESSIONAL SUMMARY, etc.)
BODY_SIZE = Pt(9.5)      # 9.5pt for body text, bullets, dates
CONTACT_SIZE = Pt(9.5)   # 9.5pt for contact info

PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11)
MARGIN = Inches(1)


# ── Cover Letter Generator ────────────────────────────────────────────────


def generate_cover_letter_docx(
    cover_letter_text: str,
    job: dict,
    output_path: Path,
    candidate_name: str = "",
    candidate_address: str = "",
) -> Path:
    """Generate a formatted .docx cover letter.

    Args:
        cover_letter_text: AI-generated cover letter body text.
        job: Job dict with title, company, location, etc.
        output_path: Directory to save the file.
        candidate_name: Candidate's full name for sign-off.
        candidate_address: Candidate's location line.

    Returns:
        Path to the generated .docx file.
    """
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # --- Header: Name + Date (tab-separated) ---
    if candidate_name:
        p = doc.add_paragraph()
        _set_tab_stops(p, right_tab=True)
        run = p.add_run(candidate_name)
        run.font.name = FONT_NAME
        run.font.size = BODY_SIZE
        date_str = f"{datetime.now().day} {datetime.now().strftime('%B %Y')}"
        run2 = p.add_run(f"\t{date_str}")
        run2.font.name = FONT_NAME
        run2.font.size = BODY_SIZE

    # Candidate address
    if candidate_address:
        p = doc.add_paragraph()
        run = p.add_run(candidate_address)
        run.font.name = FONT_NAME
        run.font.size = BODY_SIZE

    # Blank line
    doc.add_paragraph()

    # Company address block
    company = job.get("company", "")
    location = job.get("location", "")
    if company:
        p = doc.add_paragraph()
        run = p.add_run(company)
        run.font.name = FONT_NAME
        run.font.size = BODY_SIZE
    if location:
        p = doc.add_paragraph()
        run = p.add_run(location)
        run.font.name = FONT_NAME
        run.font.size = BODY_SIZE

    # Blank line before body
    doc.add_paragraph()

    # --- Cover letter body paragraphs ---
    # Split text into paragraphs, preserving structure
    paragraphs = [p.strip() for p in cover_letter_text.split("\n\n") if p.strip()]

    for para_text in paragraphs:
        # Skip redundant headers like "[DRAFT FLAG: ...]"
        if para_text.startswith("[DRAFT"):
            continue

        # Salutation / sign-off lines — plain text, no markdown
        if para_text.startswith("Dear") or para_text.startswith("Sincerely"):
            p = doc.add_paragraph()
            run = p.add_run(_strip_md(para_text))
            run.font.name = FONT_NAME
            run.font.size = BODY_SIZE
            if para_text.startswith("Dear"):
                doc.add_paragraph()
        elif _strip_md(para_text).strip() in [
            candidate_name, "Grygorii T.",
            "Grygorii T.", 'Grygorii T.',
        ]:
            p = doc.add_paragraph()
            run = p.add_run(_strip_md(para_text))
            run.font.name = FONT_NAME
            run.font.size = BODY_SIZE
        else:
            # Body paragraph — handle **bold** inline markers
            # Check if it's a bullet-style paragraph (• **Label:** text)
            sub_lines = para_text.split("\n")
            for sub_line in sub_lines:
                sub_line = sub_line.strip()
                if not sub_line:
                    continue
                p = doc.add_paragraph()
                if sub_line.startswith(("•", "-", "–")):
                    bullet_text = re.sub(r'^[•\-–]\s*', '', sub_line).strip()
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.first_line_indent = Inches(-0.25)
                    run_b = p.add_run("• ")
                    run_b.font.name = FONT_NAME
                    run_b.font.size = BODY_SIZE
                    _add_markdown_runs(p, bullet_text)
                else:
                    _add_markdown_runs(p, sub_line)
                p.paragraph_format.space_after = Pt(6)

    # Save
    output_path.mkdir(parents=True, exist_ok=True)
    company_slug = _slugify(company or "unknown")
    filename = f"GrygoriiT_CoverLetter_{company_slug}.docx"
    filepath = output_path / filename
    doc.save(str(filepath))
    logger.info(f"Cover letter saved: {filepath}")
    return filepath


# ── Resume Generator ──────────────────────────────────────────────────────


def generate_resume_docx(
    resume_text: str,
    job: dict,
    output_path: Path,
) -> Path:
    """Generate a formatted .docx resume matching the user's style.

    Parses the AI-generated resume text and applies formatting:
    - Bold name heading (13pt)
    - Contact info line (9.5pt)
    - Section headings with bottom borders (11pt bold)
    - Job titles bold with right-aligned dates
    - Bullet points as list items (9.5pt)

    Args:
        resume_text: AI-generated resume content (plain text).
        job: Job dict for filename generation.
        output_path: Directory to save the file.

    Returns:
        Path to the generated .docx file.
    """
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    # Parse resume sections
    lines = resume_text.split("\n")
    _build_resume_document(doc, lines)

    # Save
    output_path.mkdir(parents=True, exist_ok=True)
    company = job.get("company", "unknown")
    company_slug = _slugify(company)
    filename = f"GrygoriiT_Resume_{company_slug}.docx"
    filepath = output_path / filename
    doc.save(str(filepath))
    logger.info(f"Resume saved: {filepath}")
    return filepath


def _build_resume_document(doc: Document, lines: list):
    """Parse plain text resume and apply formatting to the Document."""

    # Known section headings
    SECTION_HEADINGS = {
        "PROFESSIONAL SUMMARY", "TECHNICAL SKILLS", "WORK EXPERIENCE",
        "EDUCATION", "GAPS ACKNOWLEDGED", "GROWTH AREAS",
        "SKILLS", "EXPERIENCE", "SUMMARY",
    }

    # Patterns for job title lines (bold title + tab + date)
    DATE_PATTERN = re.compile(
        r"^(.+?)\s*[\t|–—-]\s*"
        r"((?:January|February|March|April|May|June|July|August|September|October|November|December)"
        r"\s+\d{4}\s*(?:–|—|-)\s*(?:Present|"
        r"(?:January|February|March|April|May|June|July|August|September|October|November|December)"
        r"\s+\d{4}))$",
        re.IGNORECASE
    )

    # Company/location line pattern
    COMPANY_LINE_PATTERN = re.compile(
        r"^(.+?)\s*\|\s*(.+)$"
    )

    is_first_line = True
    i = 0

    while i < len(lines):
        raw_line = lines[i].strip()

        # Skip empty lines
        if not raw_line:
            i += 1
            continue

        # Strip markdown for detection, keep raw for rendering
        clean = _strip_md(raw_line)

        # --- Name heading (first non-empty line or matches known name pattern) ---
        if is_first_line and (
            "T." in clean.upper() or
            "GRISHA" in clean.upper() or
            i == 0
        ):
            is_first_line = False
            if len(clean) < 60:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(clean.upper())
                run.bold = True
                run.font.name = FONT_NAME
                run.font.size = NAME_SIZE
                i += 1
                continue

        is_first_line = False

        # --- Contact info lines (contain email, phone, linkedin, github) ---
        lower = clean.lower()
        if any(kw in lower for kw in ["email:", "mobile:", "phone:", "linkedin:", "github:", "@gmail"]):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(clean)
            run.font.name = FONT_NAME
            run.font.size = CONTACT_SIZE
            i += 1
            continue

        # --- Section headings (detect after stripping **markers**) ---
        heading_test = clean.upper().strip().rstrip(":")
        if heading_test in SECTION_HEADINGS or (clean.isupper() and len(clean) < 40) or (
            raw_line.startswith("**") and raw_line.endswith("**") and clean.isupper()
        ):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(heading_test)
            run.bold = True
            run.font.name = FONT_NAME
            run.font.size = SECTION_SIZE
            _add_bottom_border(p)
            i += 1
            continue

        # --- Job title with date (bold title + right-aligned date) ---
        date_match = DATE_PATTERN.match(clean)
        if date_match:
            title_part = date_match.group(1).strip()
            date_part = date_match.group(2).strip()
            p = doc.add_paragraph()
            _set_tab_stops(p, right_tab=True)
            run_title = p.add_run(title_part)
            run_title.bold = True
            run_title.font.name = FONT_NAME
            run_title.font.size = BODY_SIZE
            p.add_run("\t")
            run_date = p.add_run(date_part)
            run_date.font.name = FONT_NAME
            run_date.font.size = BODY_SIZE
            i += 1
            continue

        # --- Bullet points (with inline **bold** support) ---
        if raw_line.startswith(("•", "-", "–", "▪", "►", "●")) or (
            raw_line.startswith("*") and not raw_line.startswith("**")
        ):
            # Strip only the bullet char + whitespace, NOT * (which would eat **bold**)
            bullet_text = re.sub(r'^[•\-–▪►●]\s*', '', raw_line).strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            run_bullet = p.add_run("• ")
            run_bullet.font.name = FONT_NAME
            run_bullet.font.size = BODY_SIZE
            _add_markdown_runs(p, bullet_text)
            i += 1
            continue

        # --- Company / Location line (italic, smaller) ---
        company_match = COMPANY_LINE_PATTERN.match(clean)
        if company_match and i > 0:
            p = doc.add_paragraph()
            run = p.add_run(clean)
            run.font.name = FONT_NAME
            run.font.size = BODY_SIZE
            run.italic = True
            i += 1
            continue

        # --- Regular body text (with inline **bold** support) ---
        p = doc.add_paragraph()
        _add_markdown_runs(p, raw_line)
        i += 1


# ── PDF Conversion ────────────────────────────────────────────────────────


def convert_to_pdf(docx_path: Path) -> Optional[Path]:
    """Convert a .docx file to PDF using LibreOffice.

    Returns:
        Path to the PDF file, or None if conversion failed.
    """
    pdf_path = docx_path.with_suffix(".pdf")

    # Try LibreOffice
    for soffice_cmd in ["soffice", "libreoffice", r"C:\Program Files\LibreOffice\program\soffice.exe"]:
        try:
            result = subprocess.run(
                [soffice_cmd, "--headless", "--convert-to", "pdf",
                 "--outdir", str(docx_path.parent), str(docx_path)],
                capture_output=True, text=True, timeout=60,
            )
            if result.returncode == 0 and pdf_path.exists():
                logger.info(f"PDF generated: {pdf_path}")
                return pdf_path
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue

    logger.warning("LibreOffice not found — PDF conversion unavailable")
    return None


# ── Main Package Generator ────────────────────────────────────────────────


def generate_docx_package(
    resume_text: str,
    cover_letter_text: str,
    job: dict,
    output_dir: Path,
    candidate_name: str = "Grygorii T.",
    candidate_address: str = "Gaithersburg, MD 20878",
    generate_pdf: bool = False,
) -> Dict[str, Any]:
    """Generate both resume and cover letter as .docx files.

    Args:
        resume_text: AI-generated resume content.
        cover_letter_text: AI-generated cover letter content.
        job: Job dict with title, company, description, etc.
        output_dir: Directory to save files.
        candidate_name: Candidate's name for the cover letter header.
        candidate_address: Candidate's address line.
        generate_pdf: If True, also generate PDF versions.

    Returns:
        Dict with paths: resume_docx, cover_letter_docx, resume_pdf, cover_letter_pdf
    """
    result = {
        "resume_docx": None,
        "cover_letter_docx": None,
        "resume_pdf": None,
        "cover_letter_pdf": None,
    }

    try:
        # Generate resume .docx
        resume_path = generate_resume_docx(resume_text, job, output_dir)
        result["resume_docx"] = resume_path

        # Generate cover letter .docx
        cl_path = generate_cover_letter_docx(
            cover_letter_text, job, output_dir,
            candidate_name=candidate_name,
            candidate_address=candidate_address,
        )
        result["cover_letter_docx"] = cl_path

        # Optional PDF conversion
        if generate_pdf:
            if resume_path:
                result["resume_pdf"] = convert_to_pdf(resume_path)
            if cl_path:
                result["cover_letter_pdf"] = convert_to_pdf(cl_path)

    except Exception as e:
        logger.error(f"DOCX package generation failed: {e}", exc_info=True)

    return result


# ── Helpers ───────────────────────────────────────────────────────────────


def _add_markdown_runs(paragraph, text: str, font_name: str = FONT_NAME, font_size=BODY_SIZE):
    """Parse **bold** and *italic* markers in text and add properly formatted runs."""
    # Split on **bold** markers first
    parts = re.split(r'(\*\*[^*]+?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text — strip markers
            inner = part[2:-2]
            run = paragraph.add_run(inner)
            run.bold = True
        else:
            # Check for *italic* markers within normal text
            sub_parts = re.split(r'(\*[^*]+?\*)', part)
            for sub in sub_parts:
                if sub.startswith('*') and sub.endswith('*') and not sub.startswith('**'):
                    run = paragraph.add_run(sub[1:-1])
                    run.italic = True
                else:
                    if sub:
                        run = paragraph.add_run(sub)
                    else:
                        continue
                run.font.name = font_name
                run.font.size = font_size
                continue
            continue
        run.font.name = font_name
        run.font.size = font_size


def _strip_md(text: str) -> str:
    """Remove markdown bold/italic markers from text."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    return text.strip()


def _add_bottom_border(paragraph):
    """Add a bottom border to a paragraph (like the section heading lines)."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="000000"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def _set_tab_stops(paragraph, right_tab: bool = False):
    """Set tab stops on a paragraph — right-aligned tab at margin end."""
    if right_tab:
        pPr = paragraph._element.get_or_add_pPr()
        tabs = parse_xml(
            f'<w:tabs {nsdecls("w")}>'
            '  <w:tab w:val="right" w:pos="9360"/>'
            '</w:tabs>'
        )
        pPr.append(tabs)


def _slugify(text: str) -> str:
    """Convert text to filesystem-safe slug."""
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"[\s_]+", "_", text)
    return text.strip("_")[:40]
